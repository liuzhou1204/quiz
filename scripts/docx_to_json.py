#!/usr/bin/env python3
"""
docx_to_json.py - Word/PDF 题库转 JSON 工具

将 Word (.docx) 或 PDF (.pdf) 格式的题库文件转换为标准 JSON 格式，
供 Quiz Vue 系统加载使用。

使用方式:
    python docx_to_json.py -i "题库.docx" -o "public/data/pm.json" --bank "项目管理大赛题库"

依赖安装:
    pip install python-docx pdfplumber
"""

import re
import sys
import json
import argparse
from pathlib import Path

# 尝试导入 docx/pdf 解析库
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False


def read_docx(filepath):
    """读取 docx 文件，返回段落文本列表"""
    if not HAS_DOCX:
        print("错误: 需要安装 python-docx 库")
        print("  pip install python-docx")
        sys.exit(1)
    doc = Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # 也读取表格内容（有些题库把选项放在表格里）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return paragraphs


def read_pdf(filepath):
    """读取 pdf 文件，返回段落文本列表"""
    if not HAS_PDF:
        print("错误: 需要安装 pdfplumber 库")
        print("  pip install pdfplumber")
        sys.exit(1)
    paragraphs = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                for line in text.split('\n'):
                    line = line.strip()
                    if line:
                        paragraphs.append(line)
    return paragraphs


def classify_question_type(lines, q_start, answer_line_idx):
    """根据题干内容和答案判断题目类型"""
    stem = ' '.join(lines[q_start:answer_line_idx])

    # 判断题：题干通常不含选项，答案为"正确/错误/对/错"
    if answer_line_idx < len(lines):
        answer_text = lines[answer_line_idx]
        if re.search(r'答案[：:]\s*(正确|错误|对|错)', answer_text):
            return 'judge'

    # 填空题：题干含下划线或括号占位
    if re.search(r'_{2,}|\(\s*\)|（\s*）', stem):
        return 'fill'

    # 多选题：题干含"哪些""以下哪些""多项"等关键词，或答案有多个字母
    multi_keywords = ['哪些', '以下哪些', '多项', '多项选择', '多选题']
    if any(kw in stem for kw in multi_keywords):
        return 'multi'

    # 检查选项数量来判断
    option_count = 0
    for line in lines[q_start:answer_line_idx]:
        if re.match(r'^[A-F][.、)\s]', line.strip()):
            option_count += 1
    if option_count >= 5:  # 5个以上选项通常为多选
        return 'multi'

    return 'single'


def parse_options(lines, q_start, answer_line_idx):
    """从行中提取选项列表"""
    options = []
    for line in lines[q_start:answer_line_idx]:
        stripped = line.strip()
        # 匹配 A. / A、 / A) / A  开头的选项
        match = re.match(r'^([A-F])[.、)\s]\s*(.*)', stripped)
        if match:
            options.append((match.group(1), match.group(2).strip()))
    return options


def parse_answer(answer_text):
    """从答案行中提取答案"""
    # 移除 "答案:" "正确答案:" 等前缀
    answer_text = re.sub(r'^(正确)?答案\s*[：:]\s*', '', answer_text.strip())
    return answer_text.strip()


def parse_stem(lines, q_start, answer_line_idx):
    """提取题干（移除题号前缀）"""
    stem_lines = []
    for line in lines[q_start:answer_line_idx]:
        stripped = line.strip()
        # 跳过选项行
        if re.match(r'^[A-F][.、)\s]', stripped):
            continue
        # 移除题号前缀（如 "1." "1、" "第1题"）
        cleaned = re.sub(r'^\d{1,3}[.、)\s]\s*', '', stripped)
        cleaned = re.sub(r'^第?\d{1,3}题[.、:\s]*', '', cleaned)
        if cleaned:
            stem_lines.append(cleaned)
    return '\n'.join(stem_lines)


def parse_explain(lines, answer_line_idx):
    """从答案行后面提取解析"""
    explain_lines = []
    for i in range(answer_line_idx + 1, min(answer_line_idx + 3, len(lines))):
        line = lines[i].strip()
        # 跳过空行和下一题的开头
        if not line or re.match(r'^\d{1,3}[.、)\s]', line) or re.match(r'^第?\d{1,3}题', line):
            break
        # 跳过明显的非解析内容
        if line.startswith('答案') or line.startswith('正确答案'):
            continue
        explain_lines.append(line)
    return '\n'.join(explain_lines) if explain_lines else ''


def extract_category(stem, default='综合'):
    """尝试从题干推断分类"""
    category_keywords = {
        '进度管理': ['进度', '工期', '关键路径', 'CPM', 'PERT', '里程碑', '赶工', '快速跟进', '资源平衡', '资源平滑', '移峰填谷'],
        '成本管理': ['成本', '预算', '估算', '挣值', 'EVM', 'EAC', 'CPI', 'SPI', '储备', 'TCPI'],
        '质量管理': ['质量', '缺陷', 'QA', 'QC', 'PDCA', '石川图', '鱼骨图', '帕累托', '控制图', '零缺陷'],
        '范围管理': ['范围', 'WBS', '需求', '可交付成果', '范围蔓延', '确认范围', '工作包'],
        '整合管理': ['整合', '变更', 'CCB', '章程', '收尾', '项目章程', '整体变更'],
        '资源管理': ['资源', '团队', '塔克曼', '冲突', '组织结构', '矩阵', '预分派', '人员'],
        '沟通管理': ['沟通', '渠道', '会议', '相关方', '干系人', 'kickoff'],
        '风险管理': ['风险', '威胁', '机会', '应急', 'EMV', '储备', '规避', '转移', '减轻'],
        '采购管理': ['采购', '合同', '招标', '投标', '卖方', 'FFP', 'CPIF', 'T&M'],
        '相关方管理': ['相关方', '权力', '利益', '方格', '参与度'],
        '领导力': ['领导', '领导力', '领导风格', '情境领导'],
        'HSSE管理': ['HSSE', '安全', '健康', '环境', '应急'],
        '冲突管理': ['冲突'],
        '合同管理': ['合同', '要约', '承诺', '违约'],
        '团队管理': ['团队'],
        '组织管理': ['组织', 'PMO', 'IPMA'],
        '项目决策': ['决策'],
        '项目治理': ['治理'],
        '项目管理基础': ['过程组', '知识领域', 'PMBOK'],
        '项目组合': ['项目集群', '项目组合'],
        '项目评估': ['评估', '可行性', 'IRR', 'NPV'],
    }
    for category, keywords in category_keywords.items():
        for kw in keywords:
            if kw.lower() in stem.lower():
                return category
    return default


def parse_questions(lines, category_default='综合'):
    """从文本行列表中解析题目"""
    questions = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].strip()

        # 检测题目开头：数字编号 或 "第X题"
        if not re.match(r'^(\d{1,3}[.、)\s]|第?\d{1,3}题)', line):
            i += 1
            continue

        q_start = i
        i += 1

        # 找到答案行和下一题的开头
        answer_line_idx = -1
        while i < n:
            curr = lines[i].strip()
            # 检测答案行
            if re.search(r'^(正确)?答案\s*[：:]', curr):
                answer_line_idx = i
                i += 1
                break
            # 检测下一题开头（数字编号）
            if re.match(r'^\d{1,3}[.、)\s]', curr) and i > q_start + 1:
                break
            i += 1

        if answer_line_idx == -1:
            continue  # 没有找到答案行，跳过

        # 判断题目类型
        q_type = classify_question_type(lines, q_start, answer_line_idx)

        # 提取各部分
        stem = parse_stem(lines, q_start, answer_line_idx)
        answer_text = parse_answer(lines[answer_line_idx])
        explain = parse_explain(lines, answer_line_idx)
        options = parse_options(lines, q_start, answer_line_idx)
        category = extract_category(stem, category_default)

        # 构建题目对象
        question = {
            'q': stem,
            'answer': answer_text,
            'explain': explain,
            'category': category,
            'type': q_type,
        }

        if q_type in ('single', 'multi') and options:
            question['options'] = [opt[1] for opt in options]

        questions.append(question)

    return questions


def update_manifest(manifest_path, bank_id, bank_name, total, create=False):
    """更新题库清单文件"""
    if create or manifest_path.exists():
        manifest = {"banks": []}
        if manifest_path.exists():
            try:
                with open(manifest_path, 'r', encoding='utf-8') as f:
                    manifest = json.load(f)
            except (json.JSONDecodeError, IOError):
                pass

        # 查找并更新或新增
        banks = manifest.get('banks', [])
        found = False
        for b in banks:
            if b.get('id') == bank_id:
                b['name'] = bank_name
                b['total'] = total
                b['file'] = f"{bank_id}.json"
                found = True
                break
        if not found:
            banks.append({
                'id': bank_id,
                'name': bank_name,
                'file': f"{bank_id}.json",
                'total': total
            })

        manifest['banks'] = banks
        with open(manifest_path, 'w', encoding='utf-8') as f:
            json.dump(manifest, f, ensure_ascii=False, indent=2)
        print(f"  已更新题库清单: {manifest_path}")


def convert(input_path, output_path, bank_name, category_default='综合'):
    """主转换函数"""
    input_file = Path(input_path)
    if not input_file.exists():
        print(f"错误: 文件不存在 - {input_path}")
        sys.exit(1)

    # 读取文件
    ext = input_file.suffix.lower()
    if ext == '.docx':
        print(f"正在读取 Word 文件: {input_path}")
        lines = read_docx(input_path)
    elif ext == '.pdf':
        print(f"正在读取 PDF 文件: {input_path}")
        lines = read_pdf(input_path)
    else:
        print(f"错误: 不支持的文件格式 - {ext}（仅支持 .docx 和 .pdf）")
        sys.exit(1)

    print(f"  共读取 {len(lines)} 行文本")

    # 解析题目
    questions = parse_questions(lines, category_default)
    print(f"  成功解析 {len(questions)} 道题目")

    if not questions:
        print("警告: 未解析到任何题目，请检查文件格式")
        print("  题目格式要求：")
        print("  1. 每题以数字编号开头（如 '1.' 或 '1、'）")
        print("  2. 答案行以 '答案:' 或 '正确答案:' 开头")
        print("  3. 选项以 A. B. C. D. 开头")
        sys.exit(1)

    # 统计
    type_counts = {}
    categories = set()
    for q in questions:
        t = q['type']
        type_counts[t] = type_counts.get(t, 0) + 1
        categories.add(q['category'])

    print(f"\n  题型分布: {json.dumps(type_counts, ensure_ascii=False)}")
    print(f"  分类: {', '.join(sorted(categories))}")

    # 检查空解析
    no_explain = sum(1 for q in questions if not q.get('explain'))
    if no_explain:
        print(f"  警告: {no_explain} 道题目缺少解析")

    # 构建 JSON
    data = {
        'meta': {
            'name': bank_name,
            'version': '1.0.0',
            'total': len(questions),
            'categories': sorted(categories),
            'typeCounts': type_counts,
        },
        'questions': questions
    }

    # 写入文件
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print(f"\n  已写入: {output_file}")

    # 自动更新 manifest.json
    manifest_path = output_file.parent / 'manifest.json'
    if manifest_path.exists():
        update_manifest(manifest_path, output_file.stem, bank_name, len(questions))
    elif output_file.parent.name == 'data':
        # 如果在 public/data/ 下，自动创建/更新 manifest
        manifest_path = output_file.parent / 'manifest.json'
        update_manifest(manifest_path, output_file.stem, bank_name, len(questions), create=True)


def main():
    parser = argparse.ArgumentParser(
        description='Word/PDF 题库转 JSON 工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python docx_to_json.py -i 题库.docx -o public/data/new.json --bank "新题库"
  python docx_to_json.py -i 题库.pdf -o public/data/pmp.json --bank "PMP题库"

支持的题目格式:
  单选题:
    1. 题干内容
    A. 选项A
    B. 选项B
    C. 选项C
    D. 选项D
    答案: B

  多选题:
    2. 以下哪些是正确的
    A. 选项A
    B. 选项B
    C. 选项C
    答案: A,B,C  或  答案: ABC

  判断题:
    3. 题干内容
    答案: 正确  或  答案: 对

  填空题:
    4. 题干内容____
    答案: 参考答案文本
        """
    )
    parser.add_argument('-i', '--input', required=True, help='输入文件路径 (.docx 或 .pdf)')
    parser.add_argument('-o', '--output', required=True, help='输出 JSON 文件路径')
    parser.add_argument('--bank', default='未命名题库', help='题库名称（默认: 未命名题库）')
    parser.add_argument('--category', default='综合', help='默认分类（无法识别时使用，默认: 综合）')

    args = parser.parse_args()
    convert(args.input, args.output, args.bank, args.category)


if __name__ == '__main__':
    main()
